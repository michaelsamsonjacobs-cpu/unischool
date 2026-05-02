import sys
import os
import json
import re
import csv

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import docx as python_docx
except ImportError:
    python_docx = None

def shred_single_file(filepath):
    """Extract text from a single file. Returns (text, error)."""
    if not os.path.exists(filepath):
        return "", f"File not found: {filepath}"

    extension = os.path.splitext(filepath)[1].lower()
    raw_text = ""
    basename = os.path.basename(filepath)

    if extension == '.pdf':
        if not PyPDF2:
            return "", "PyPDF2 module not installed for PDF parsing."
        try:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        raw_text += extracted + "\n"
        except Exception as e:
            return "", f"Failed to parse PDF '{basename}': {str(e)}"

    elif extension in ['.xlsx', '.xls']:
        if not openpyxl:
            return "", "openpyxl module not installed for Excel parsing."
        try:
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                raw_text += f"\n=== Sheet: {sheet_name} ===\n"
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    raw_text += " | ".join(cells) + "\n"
            wb.close()
        except Exception as e:
            return "", f"Failed to parse Excel '{basename}': {str(e)}"

    elif extension == '.csv':
        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f)
                for row in reader:
                    raw_text += " | ".join(row) + "\n"
        except Exception as e:
            return "", f"Failed to parse CSV '{basename}': {str(e)}"

    elif extension in ['.docx', '.doc']:
        if not python_docx:
            return "", "python-docx module not installed for Word document parsing."
        try:
            doc = python_docx.Document(filepath)
            # Extract all paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    raw_text += para.text + "\n"
            # Extract table contents (many BAA/SOW docs use tables for requirements)
            for table in doc.tables:
                raw_text += "\n--- TABLE ---\n"
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    raw_text += " | ".join(cells) + "\n"
        except Exception as e:
            return "", f"Failed to parse Word document '{basename}': {str(e)}"

    elif extension in ['.md', '.txt']:
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                raw_text = f.read()
        except Exception as e:
            return "", f"Failed to read text file '{basename}': {str(e)}"
    elif extension == '.zip':
        import zipfile
        import tempfile
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                with zipfile.ZipFile(filepath, 'r') as zip_ref:
                    zip_ref.extractall(tmpdir)
                # Recursively shred extracted files
                extracted_files = [os.path.join(dp, f) for dp, dn, filenames in os.walk(tmpdir) for f in filenames]
                for f in extracted_files:
                    t, e = shred_single_file(f)
                    if not e: raw_text += t + "\n"
        except Exception as e:
            return "", f"Failed to parse ZIP '{basename}': {str(e)}"
    else:
        return "", f"Unsupported file type: {extension}"

    return raw_text, None


def shred_documents(filepaths):
    """Shred one or more files and combine the output."""
    combined_text = ""
    errors = []
    files_processed = []

    for filepath in filepaths:
        basename = os.path.basename(filepath)
        text, err = shred_single_file(filepath)
        if err:
            errors.append(err)
        else:
            combined_text += f"\n\n{'='*60}\nDOCUMENT: {basename}\n{'='*60}\n\n{text}"
            files_processed.append(basename)

    if not combined_text and errors:
        print(json.dumps({"error": "; ".join(errors)}))
        return

    # RTM (Requirements Traceability Matrix) Extraction
    # Scan explicitly for strict legal "shall" or "must" modifiers
    rtm_shalls = []
    sentences = re.split(r'(?<=[.!?]) +', combined_text.replace('\n', ' '))
    for idx, sentence in enumerate(sentences):
        sentence_lower = sentence.lower()
        if ' shall ' in sentence_lower or ' must ' in sentence_lower:
            rtm_shalls.append({
                "line_id": f"RFP-REQ-{idx+1:04d}",
                "directive": sentence.strip()
            })

    print(json.dumps({
        "status": "success",
        "doc_length": len(combined_text),
        "files_processed": files_processed,
        "files_count": len(files_processed),
        "raw_content": combined_text,
        "rtm_matrix": rtm_shalls,
        "warnings": errors if errors else []
    }))


if __name__ == "__main__":
    if len(sys.argv) > 1:
        shred_documents(sys.argv[1:])
    else:
        print(json.dumps({"error": "No filepath(s) provided."}))
